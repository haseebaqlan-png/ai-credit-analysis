# AI Credit Analysis MVP 1.0 — Mobile Upload Edition
# Generated as a single-file application for easy GitHub upload from mobile.

import os
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker, declarative_base
DATABASE_URL=os.getenv("DATABASE_URL","sqlite:///./app.db")
connect_args={"check_same_thread":False} if DATABASE_URL.startswith("sqlite") else {}
engine=create_engine(DATABASE_URL,future=True,pool_pre_ping=True,connect_args=connect_args)
SessionLocal=sessionmaker(bind=engine,autoflush=False,autocommit=False,future=True)
Base=declarative_base()
def get_db():
    db=SessionLocal()
    try: yield db
    finally: db.close()

from sqlalchemy import Column,Integer,String,Float,Boolean,DateTime,ForeignKey,Text
from sqlalchemy.orm import relationship
from datetime import datetime
class User(Base):
    __tablename__='users';id=Column(Integer,primary_key=True);email=Column(String(190),unique=True,index=True,nullable=False);password_hash=Column(String(255),nullable=False);full_name=Column(String(160),default='');created_at=Column(DateTime,default=datetime.utcnow)
class Client(Base):
    __tablename__='clients';id=Column(Integer,primary_key=True);user_id=Column(Integer,ForeignKey('users.id'),nullable=False,index=True);name=Column(String(220),nullable=False);legal_form=Column(String(120),default='');sector=Column(String(160),default='');registration_no=Column(String(120),default='');city=Column(String(120),default='');notes=Column(Text,default='');created_at=Column(DateTime,default=datetime.utcnow);applications=relationship('CreditApplication',back_populates='client',cascade='all, delete-orphan')
class CreditApplication(Base):
    __tablename__='credit_applications';id=Column(Integer,primary_key=True);client_id=Column(Integer,ForeignKey('clients.id'),nullable=False,index=True);title=Column(String(220),default='طلب تمويل');facility_type=Column(String(120),default='');purpose=Column(Text,default='');requested_amount=Column(Float,default=0);currency=Column(String(20),default='YER');tenor_months=Column(Float,default=12);annual_rate=Column(Float,default=0);existing_annual_debt_service=Column(Float,default=0);collateral_value=Column(Float,default=0);repayment_source=Column(Text,default='');qualitative_score=Column(Float,default=50);days_past_due=Column(Float,default=0);default_flag=Column(Boolean,default=False);status=Column(String(40),default='draft');created_at=Column(DateTime,default=datetime.utcnow);client=relationship('Client',back_populates='applications');documents=relationship('Document',back_populates='application',cascade='all, delete-orphan');values=relationship('FinancialValue',back_populates='application',cascade='all, delete-orphan');costs=relationship('ApiCost',back_populates='application',cascade='all, delete-orphan')
class Document(Base):
    __tablename__='documents';id=Column(Integer,primary_key=True);application_id=Column(Integer,ForeignKey('credit_applications.id'),nullable=False,index=True);filename=Column(String(255),nullable=False);stored_path=Column(String(500),nullable=False);file_type=Column(String(40),default='');size_bytes=Column(Integer,default=0);extraction_status=Column(String(40),default='pending');extracted_count=Column(Integer,default=0);created_at=Column(DateTime,default=datetime.utcnow);application=relationship('CreditApplication',back_populates='documents')
class FinancialValue(Base):
    __tablename__='financial_values';id=Column(Integer,primary_key=True);application_id=Column(Integer,ForeignKey('credit_applications.id'),nullable=False,index=True);document_id=Column(Integer,ForeignKey('documents.id'),nullable=True,index=True);period=Column(String(40),nullable=False);statement=Column(String(50),default='financial');metric=Column(String(80),nullable=False);value=Column(Float,default=0);source_location=Column(String(255),default='');confidence=Column(Float,default=0);verified=Column(Boolean,default=False);reviewer_note=Column(Text,default='');application=relationship('CreditApplication',back_populates='values')
class ApiCost(Base):
    __tablename__='api_costs';id=Column(Integer,primary_key=True);application_id=Column(Integer,ForeignKey('credit_applications.id'),nullable=True,index=True);provider=Column(String(80),default='local');model=Column(String(120),default='rules-v1');operation=Column(String(120),default='');input_units=Column(Integer,default=0);output_units=Column(Integer,default=0);cost_usd=Column(Float,default=0);created_at=Column(DateTime,default=datetime.utcnow);application=relationship('CreditApplication',back_populates='costs')

import bcrypt
def hash_password(password:str)->str:return bcrypt.hashpw(password.encode(),bcrypt.gensalt()).decode()
def verify_password(password:str,password_hash:str)->bool:
    try:return bcrypt.checkpw(password.encode(),password_hash.encode())
    except:return False

from openpyxl import load_workbook
from pypdf import PdfReader
from docx import Document as DocxDocument
import io,re,csv,os
ALIASES={
'revenue':['revenue','sales','turnover','الإيرادات','الايرادات','المبيعات'],'cogs':['cost of goods sold','cost of sales','cogs','تكلفة المبيعات','تكلفة البضاعة المباعة'],'gross_profit':['gross profit','مجمل الربح','إجمالي الربح','اجمالي الربح'],'ebitda':['ebitda','الربح قبل الفوائد والضرائب والاستهلاك والإطفاء'],'operating_profit':['operating profit','ebit','الربح التشغيلي'],'net_profit':['net profit','net income','صافي الربح','صافي الدخل'],'cash':['cash','cash and cash equivalents','النقد','النقدية','النقد وما في حكمه'],'receivables':['receivables','accounts receivable','trade receivables','الذمم المدينة','المدينون'],'inventory':['inventory','stocks','المخزون','المخزون السلعي'],'current_assets':['current assets','الأصول المتداولة','الاصول المتداولة'],'fixed_assets':['property plant equipment','ppe','fixed assets','الأصول الثابتة','الاصول الثابتة'],'total_assets':['total assets','إجمالي الأصول','اجمالي الاصول'],'payables':['payables','accounts payable','trade payables','الذمم الدائنة','الدائنون'],'short_term_debt':['short term debt','current borrowings','قروض قصيرة الأجل','تمويلات قصيرة الأجل'],'current_liabilities':['current liabilities','الالتزامات المتداولة','الخصوم المتداولة'],'long_term_debt':['long term debt','non-current borrowings','قروض طويلة الأجل','تمويلات طويلة الأجل'],'total_debt':['total debt','borrowings','loans','إجمالي الدين','اجمالي الدين','القروض'],'total_liabilities':['total liabilities','إجمالي الالتزامات','اجمالي الالتزامات','إجمالي الخصوم','اجمالي الخصوم'],'equity':['equity','shareholders equity','حقوق الملكية','حقوق المساهمين'],'cfo':['cash flow from operations','operating cash flow','net cash from operating activities','التدفق النقدي التشغيلي','صافي النقد من الأنشطة التشغيلية'],'capex':['capital expenditure','capex','شراء أصول ثابتة','نفقات رأسمالية'],'interest_expense':['interest expense','finance cost','finance costs','مصروف الفوائد','تكلفة التمويل','تكاليف التمويل']}
def norm(s):
 s=str(s or '').strip().lower()
 for a,b in [('إ','ا'),('أ','ا'),('آ','ا'),('ى','ي'),('ة','ه')]:s=s.replace(a,b)
 return re.sub(r'\s+',' ',s)
NAL={k:[norm(x) for x in v] for k,v in ALIASES.items()}
def metric(label):
 z=norm(label)
 for k,als in NAL.items():
  for a in als:
   if z==a or (len(a)>=5 and a in z):return k
def num(v):
 if v is None:return None
 if isinstance(v,(int,float)):return float(v)
 s=str(v).strip()
 if not s:return None
 neg=s.startswith('(') and s.endswith(')');s=s.strip('()').replace(',','').replace('٬','').replace(' ','').replace('−','-')
 if not re.fullmatch(r'[-+]?\d*\.?\d+',s):return None
 x=float(s);return -x if neg else x
def yearcols(row):
 out=[]
 for i,v in enumerate(row):
  m=re.search(r'(20\d{2}|19\d{2})',str(v or ''))
  if m:out.append((i,m.group(1)))
 return out
def extract_xlsx(data,name):
 wb=load_workbook(io.BytesIO(data),data_only=True,read_only=True);values=[]
 for ws in wb.worksheets[:40]:
  rows=[list(r[:100]) for r in ws.iter_rows(values_only=True)];hdr=[]
  for r in rows[:25]:
   y=yearcols(r)
   if y:hdr=y;break
  if not hdr:continue
  for ri,row in enumerate(rows,1):
   key=None
   for v in row[:15]:
    if isinstance(v,str) and v.strip():
     key=metric(v)
     if key:break
   if not key:continue
   for ci,yr in hdr:
    if ci<len(row):
     val=num(row[ci])
     if val is not None:values.append(dict(period=yr,metric=key,value=val,location=f'{ws.title}!R{ri}C{ci+1}',confidence=.99))
 return values
def _text_extract(text,name):
 values=[];page='text'
 for line in [x.strip() for x in text.splitlines() if x.strip()]:
  if line.startswith('--- PAGE'):page=line
  key=metric(line)
  if not key:continue
  yrs=re.findall(r'(20\d{2}|19\d{2})',line);nums=[]
  for s in re.findall(r'\(?-?\d[\d,٬]*(?:\.\d+)?\)?',line):
   q=s.replace(',','').replace('٬','')
   if re.fullmatch(r'20\d{2}|19\d{2}',q):continue
   try:nums.append(float(q.strip('()'))*(-1 if s.startswith('(') else 1))
   except:pass
  if nums:values.append(dict(period=yrs[-1] if yrs else 'Latest',metric=key,value=nums[-1],location=page,confidence=.68))
 return values
def extract_pdf(data,name):
 r=PdfReader(io.BytesIO(data));out=[]
 for i,p in enumerate(r.pages[:120],1):
  t=p.extract_text() or ''
  if t.strip():out.append(f'--- PAGE {i} ---\n{t}')
 return _text_extract('\n'.join(out),name)
def extract_docx(data,name):
 d=DocxDocument(io.BytesIO(data));out=[p.text for p in d.paragraphs if p.text.strip()]
 for t in d.tables:
  for r in t.rows:out.append(' | '.join(c.text for c in r.cells))
 return _text_extract('\n'.join(out),name)
def extract_csv(data,name):
 text=''
 for enc in ('utf-8-sig','utf-8','cp1256','latin-1'):
  try:text=data.decode(enc);break
  except:pass
 rows=list(csv.reader(io.StringIO(text)));hdr=[];values=[]
 for r in rows[:25]:
  y=yearcols(r)
  if y:hdr=y;break
 for ri,row in enumerate(rows,1):
  key=None
  for v in row[:10]:
   key=metric(v)
   if key:break
  if not key:continue
  for ci,yr in hdr:
   if ci<len(row):
    val=num(row[ci])
    if val is not None:values.append(dict(period=yr,metric=key,value=val,location=f'R{ri}C{ci+1}',confidence=.98))
 return values
def extract_file(filename,data):
 ext=os.path.splitext(filename.lower())[1]
 if ext in ('.xlsx','.xlsm'):return extract_xlsx(data,filename)
 if ext=='.csv':return extract_csv(data,filename)
 if ext=='.pdf':return extract_pdf(data,filename)
 if ext=='.docx':return extract_docx(data,filename)
 if ext=='.txt':
  text=''
  for enc in ('utf-8-sig','utf-8','cp1256','latin-1'):
   try:text=data.decode(enc);break
   except:pass
  return _text_extract(text,filename)
 return []

def n(x,d=0):
 try:return float(x)
 except:return d
def dv(a,b,d=0):
 b=n(b);return d if b==0 else n(a)/b
def clamp(x,a=0,b=100):return max(a,min(b,n(x)))
def periods_from_values(values,verified_only=False):
 p={}
 for v in values:
  if verified_only and not v.verified:continue
  p.setdefault(v.period,{})[v.metric]=v.value
 out=[]
 for yr in sorted(p):
  x={'period':yr};x.update(p[yr])
  if not x.get('gross_profit') and x.get('revenue') and x.get('cogs'):x['gross_profit']=x['revenue']-x['cogs']
  if not x.get('total_debt'):x['total_debt']=n(x.get('short_term_debt'))+n(x.get('long_term_debt'))
  out.append(x)
 return out
def ratios(p):
 rev=n(p.get('revenue'));cogs=n(p.get('cogs'));gp=n(p.get('gross_profit'));eb=n(p.get('ebitda'));op=n(p.get('operating_profit'));ni=n(p.get('net_profit'));cash=n(p.get('cash'));ar=n(p.get('receivables'));inv=n(p.get('inventory'));ca=n(p.get('current_assets'));ta=n(p.get('total_assets'));ap=n(p.get('payables'));cl=n(p.get('current_liabilities'));debt=n(p.get('total_debt'));tl=n(p.get('total_liabilities'));eq=n(p.get('equity'));cfo=n(p.get('cfo'));capex=n(p.get('capex'));interest=n(p.get('interest_expense'));wc=ca-cl
 return {'gross_margin':dv(gp,rev),'ebitda_margin':dv(eb,rev),'operating_margin':dv(op,rev),'net_margin':dv(ni,rev),'roa':dv(ni,ta),'roe':dv(ni,eq),'asset_turnover':dv(rev,ta),'current_ratio':dv(ca,cl,99),'quick_ratio':dv(ca-inv,cl,99),'cash_ratio':dv(cash,cl,99),'working_capital':wc,'debt_equity':dv(debt,eq,99),'debt_assets':dv(debt,ta),'liabilities_assets':dv(tl,ta),'equity_ratio':dv(eq,ta),'interest_coverage':dv(eb,interest,99 if eb>0 else 0),'cfo_margin':dv(cfo,rev),'cfo_debt':dv(cfo,debt,99 if cfo>0 else 0),'free_cash_flow':cfo-capex,'dso':dv(ar,rev)*365 if rev else 0,'dio':dv(inv,cogs)*365 if cogs else 0,'dpo':dv(ap,cogs)*365 if cogs else 0,'cash_conversion_cycle':(dv(ar,rev)*365 if rev else 0)+(dv(inv,cogs)*365 if cogs else 0)-(dv(ap,cogs)*365 if cogs else 0),'inventory_turnover':dv(cogs,inv),'receivables_turnover':dv(rev,ar),'payables_turnover':dv(cogs,ap),'revenue':rev,'ebitda':eb,'net_profit':ni,'cfo':cfo,'total_assets':ta,'total_debt':debt,'equity':eq}
def trend(periods,key):
 vals=[n(x.get(key)) for x in periods if n(x.get(key))!=0]
 return 0 if len(vals)<2 else dv(vals[-1]-vals[0],abs(vals[0]))
def consistency(periods):
 issues=[]
 for p in periods:
  ta=n(p.get('total_assets'));tl=n(p.get('total_liabilities'));eq=n(p.get('equity'))
  if ta and (tl or eq):
   gap=abs(ta-(tl+eq))/max(abs(ta),1)
   if gap>.03:issues.append(f"{p['period']}: الأصول لا تساوي الالتزامات + حقوق الملكية (فرق {gap*100:.1f}%).")
  r=n(p.get('revenue'));c=n(p.get('cogs'));gp=n(p.get('gross_profit'))
  if r and c and gp and abs(gp-(r-c))/max(abs(gp),1)>.05:issues.append(f"{p['period']}: مجمل الربح لا يتطابق مع الإيرادات ناقص التكلفة.")
 return issues
def score_application(periods,app):
 p=periods[-1] if periods else {};r=ratios(p);rate=n(app.annual_rate);rate=rate/100 if rate>1 else rate;years=max(n(app.tenor_months)/12,1);debt_service=n(app.existing_annual_debt_service)+dv(app.requested_amount,years)+n(app.requested_amount)*rate;dscr=dv(p.get('cfo'),debt_service,99 if n(p.get('cfo'))>0 else 0)
 def lin(v,bad,good,higher=True):
  z=(v-bad)/(good-bad)*100 if higher else (bad-v)/(bad-good)*100
  return clamp(z)
 finance=lin(r['current_ratio'],.8,1.8)*.10+lin(r['quick_ratio'],.4,1.1)*.08+lin(r['debt_equity'],4,.8,False)*.15+lin(r['interest_coverage'],1,5)*.12+lin(dscr,.8,1.75)*.27+lin(r['ebitda_margin'],.03,.20)*.10+(100 if r['cfo']>0 else 10)*.10+lin(r['cash_conversion_cycle'],150,30,False)*.08
 score=round(finance*.75+clamp(app.qualitative_score)*.25,1)
 grade,risk=(('1','منخفضة جداً') if score>=90 else ('2','منخفضة') if score>=82 else ('3','متوسطة-منخفضة') if score>=74 else ('4','متوسطة') if score>=66 else ('5','متوسطة-مرتفعة') if score>=58 else ('6','مرتفعة') if score>=50 else ('7','مرتفعة جداً') if score>=40 else ('8','حرجة'))
 cash_limit=max(dv(r['cfo'],1.25)-n(app.existing_annual_debt_service),0)/((1/years)+rate) if ((1/years)+rate)>0 else 0;leverage_limit=max(3*n(p.get('equity'))-n(p.get('total_debt')),0);collateral_limit=dv(app.collateral_value,1.2) if n(app.collateral_value)>0 else n(app.requested_amount);indicative_limit=max(0,min(n(app.requested_amount),cash_limit,leverage_limit,collateral_limit))
 decision='لا يوصى بالموافقة بالشكل الحالي' if app.default_flag or n(app.days_past_due)>=90 or dscr<1 else 'مؤهل للموافقة المشروطة' if score>=74 and dscr>=1.25 else 'مراجعة وإعادة هيكلة قبل العرض على اللجنة'
 stresses=[{'name':'الأساسي','dscr':dscr},{'name':'ضغط متوسط','dscr':dv(r['cfo']*.75,debt_service)},{'name':'ضغط شديد','dscr':dv(r['cfo']*.55,debt_service)}]
 return {'score':score,'grade':grade,'risk':risk,'decision':decision,'ratios':r,'dscr':dscr,'cash_limit':cash_limit,'leverage_limit':leverage_limit,'collateral_limit':collateral_limit,'indicative_limit':indicative_limit,'stresses':stresses,'revenue_trend':trend(periods,'revenue'),'ebitda_trend':trend(periods,'ebitda'),'cfo_trend':trend(periods,'cfo'),'consistency':consistency(periods)}
def local_risk_analysis(periods,app,result):
 r=result['ratios'];strengths=[];risks=[];mitigants=[];conditions=[];questions=[]
 if result['dscr']>=1.25:strengths.append(f"قدرة خدمة الدين مقبولة مبدئياً؛ DSCR = {result['dscr']:.2f}x.")
 else:risks.append(f"قدرة خدمة الدين ضعيفة/حساسة؛ DSCR = {result['dscr']:.2f}x.")
 if r['current_ratio']>=1.2:strengths.append(f"السيولة الجارية عند {r['current_ratio']:.2f}x.")
 else:risks.append(f"السيولة الجارية عند {r['current_ratio']:.2f}x أقل من المستوى المريح.")
 if r['debt_equity']<=2:strengths.append(f"الرفع المالي عند {r['debt_equity']:.2f}x.")
 else:risks.append(f"الرفع المالي مرتفع عند {r['debt_equity']:.2f}x.")
 if r['cfo']>0:strengths.append('التدفق النقدي التشغيلي موجب في آخر فترة.')
 else:risks.append('التدفق النقدي التشغيلي غير موجب في آخر فترة.')
 if result['revenue_trend']<-.1:risks.append('اتجاه الإيرادات سلبي بصورة جوهرية عبر الفترات المتاحة.')
 if result['cfo_trend']<-.2:risks.append('تراجع جوهري في التدفق النقدي التشغيلي.')
 risks+=result['consistency'];mitigants+=['ربط السداد بالتدفقات التشغيلية الفعلية ومراقبة التحصيل.','استكمال الوثائق والقيم غير المعتمدة قبل قرار اللجنة.'];conditions+=['اعتماد جميع القيم المالية الجوهرية في شاشة التحقق.','استكمال KYC/AML والوثائق القانونية وفق سياسة البنك.','الحصول على مستندات الضمانات وتقييمها قانونياً وفنياً عند الانطباق.'];questions+=['ما مصدر السداد الأساسي والثانوي؟','ما أسباب التغير في رأس المال العامل؟','هل توجد التزامات خارج الميزانية أو ضمانات للغير؟','ما مدى اعتماد الإيرادات على عميل/مورد رئيسي؟']
 return {'strengths':strengths,'risks':risks,'mitigants':mitigants,'conditions':conditions,'questions':questions}

from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate,Paragraph,Spacer,Table,TableStyle
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.styles import getSampleStyleSheet,ParagraphStyle
from reportlab.lib.enums import TA_RIGHT
from arabic_reshaper import reshape
from bidi.algorithm import get_display
import os,io
FONT='/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'
if os.path.exists(FONT):pdfmetrics.registerFont(TTFont('Arabic',FONT));FONT_NAME='Arabic'
else:FONT_NAME='Helvetica'
def ar(s):
 try:return get_display(reshape(str(s)))
 except:return str(s)
def build_pdf(client,app,periods,result,analysis):
 buf=io.BytesIO();doc=SimpleDocTemplate(buf,pagesize=A4,rightMargin=32,leftMargin=32,topMargin=32,bottomMargin=32);styles=getSampleStyleSheet();normal=ParagraphStyle('ar',parent=styles['BodyText'],fontName=FONT_NAME,fontSize=9,leading=14,alignment=TA_RIGHT);title=ParagraphStyle('art',parent=normal,fontSize=17,leading=24);h2=ParagraphStyle('arh',parent=normal,fontSize=12,leading=18,spaceBefore=10,spaceAfter=6)
 story=[Paragraph(ar('مذكرة ائتمانية — مسودة للمراجعة البشرية'),title),Spacer(1,8),Paragraph(ar(f'العميل: {client.name}'),normal),Paragraph(ar(f'الطلب: {app.requested_amount:,.0f} {app.currency}'),normal),Paragraph(ar(f"التقييم: {result['score']:.0f}/100 | المخاطر: {result['grade']} — {result['risk']}"),normal),Paragraph(ar(f"التوصية الإرشادية: {result['decision']}"),normal),Spacer(1,10)]
 story.append(Paragraph(ar('المؤشرات الرئيسية'),h2));rr=result['ratios'];data=[[ar('المؤشر'),ar('القيمة')],[ar('DSCR'),f"{result['dscr']:.2f}x"],[ar('Current Ratio'),f"{rr['current_ratio']:.2f}x"],[ar('Quick Ratio'),f"{rr['quick_ratio']:.2f}x"],[ar('Debt/Equity'),f"{rr['debt_equity']:.2f}x"],[ar('Interest Coverage'),f"{rr['interest_coverage']:.2f}x"],[ar('EBITDA Margin'),f"{rr['ebitda_margin']*100:.1f}%"],[ar('CFO Margin'),f"{rr['cfo_margin']*100:.1f}%"]];t=Table(data,colWidths=[260,180],hAlign='RIGHT');t.setStyle(TableStyle([('FONT',(0,0),(-1,-1),FONT_NAME),('GRID',(0,0),(-1,-1),.4,colors.lightgrey),('BACKGROUND',(0,0),(-1,0),colors.HexColor('#E9F0F6')),('ALIGN',(0,0),(-1,-1),'RIGHT'),('PADDING',(0,0),(-1,-1),6)]));story.append(t)
 for label,key in [('نقاط القوة','strengths'),('المخاطر','risks'),('المخففات','mitigants'),('الشروط المقترحة','conditions'),('أسئلة اللجنة','questions')]:
  story.append(Paragraph(ar(label),h2))
  for item in analysis.get(key,[]):story.append(Paragraph(ar('• '+item),normal))
 story.append(Spacer(1,8));story.append(Paragraph(ar('تنبيه: هذه المذكرة أداة دعم قرار ولا تمثل موافقة ائتمانية آلية.'),normal));doc.build(story);buf.seek(0);return buf.getvalue()


# Embedded UI assets: generated automatically at runtime; no folders need uploading.

EMBEDDED_CSS = '*{box-sizing:border-box}body{margin:0;background:#eef3f8;color:#13243c;font-family:system-ui,-apple-system,"Segoe UI",Tahoma,Arial}a{color:#0b6b78;text-decoration:none}header{background:linear-gradient(120deg,#071a34,#0d5272);color:#fff}.top,.wrap{max-width:1280px;margin:auto}.top{padding:16px;display:flex;justify-content:space-between;align-items:center}.top b{font-size:22px}.top small{display:block;color:#c6d7e5}.top nav{display:flex;gap:16px}.top nav a{color:#fff}.wrap{padding:18px 12px 60px}.card{background:#fff;border:1px solid #dbe5ee;border-radius:20px;box-shadow:0 10px 30px #0b244012;margin-bottom:15px}.sec{padding:22px}.hero{padding:24px;background:linear-gradient(135deg,#fff 62%,#e7faf7);display:flex;justify-content:space-between;gap:14px;align-items:center}.eyebrow{font-size:12px;color:#0b8a7f;font-weight:900}h1{margin:7px 0;font-size:30px}h2{margin:0 0 14px}h3{margin:18px 0 8px}p{line-height:1.8;color:#657386}.btn,button{display:inline-block;background:#087d73;color:white;border:0;border-radius:11px;padding:12px 16px;font-weight:800;cursor:pointer}.ghost{background:#fff;color:#087d73;border:1px solid #99b9bb}.badge{display:inline-block;background:#edf4f8;border-radius:999px;padding:6px 10px}.notice{padding:12px;border-radius:12px;background:#fff8e6;border:1px solid #f0d68a;margin-bottom:12px}.kpis{display:grid;grid-template-columns:repeat(4,1fr);gap:10px;margin-bottom:15px}.kpi,.ratio{background:#fff;border:1px solid #dbe5ee;border-radius:15px;padding:14px}.kpi span,.ratio span{font-size:12px;color:#697688}.kpi b,.ratio b{display:block;font-size:22px;margin-top:5px}.kpi small{display:block;color:#718096}.grid{display:grid;grid-template-columns:repeat(3,1fr);gap:12px}.client{padding:20px}.client b,.client span,.client small{display:block}.client b{font-size:20px}.client span{margin-top:8px}.client small{color:#788797;margin-top:5px}.heading{display:flex;justify-content:space-between;align-items:center;margin-bottom:14px}.formgrid{display:grid;grid-template-columns:repeat(3,1fr);gap:12px}.span2{grid-column:span 2}.span3{grid-column:span 3}label{display:block;font-size:12px;color:#617084;margin-bottom:5px}input,textarea,select{width:100%;padding:11px;border:1px solid #cbd7e2;border-radius:10px;background:#fff}textarea{min-height:90px}.auth{max-width:460px;margin:60px auto;padding:26px}.auth input{margin-bottom:12px}.table{overflow:auto}table{width:100%;border-collapse:collapse;min-width:720px}th,td{padding:10px;border-bottom:1px solid #e5ebf1;text-align:right;font-size:13px}th{background:#f8fafc}.upload{display:grid;grid-template-columns:1fr auto;gap:10px;margin-bottom:16px}.tabs{display:flex;gap:8px;overflow:auto;margin:0 0 14px}.tabs a{background:#fff;border:1px solid #dbe5ee;border-radius:10px;padding:10px 14px;white-space:nowrap}.ratio-grid{display:grid;grid-template-columns:repeat(5,1fr);gap:8px}.ratio b{font-size:17px}.chart-grid{display:grid;grid-template-columns:1fr 1fr;gap:10px}.chart{border:1px solid #dbe5ee;border-radius:14px;padding:12px;overflow:auto}.chart svg{width:100%;min-width:420px}.chart path{fill:none;stroke:#0f766e;stroke-width:4}.chart circle{fill:#0f766e}.chart text{font-size:10px;fill:#657386}.twocol{display:grid;grid-template-columns:1fr 1fr;gap:20px}ul{line-height:1.9}.decision{padding:14px;border-radius:14px;font-weight:900}.green{background:#dcfce7;color:#14532d}.amber{background:#fef3c7;color:#7c4a05}.red{background:#fee2e2;color:#8b1b1b}.memo{max-width:900px;margin-left:auto;margin-right:auto}.memo .disclaimer{border-top:1px solid #dde4eb;padding-top:12px}.check{width:auto}@media(max-width:900px){.ratio-grid{grid-template-columns:repeat(3,1fr)}.grid{grid-template-columns:1fr 1fr}}@media(max-width:650px){.top{align-items:flex-start}.top nav{font-size:12px;gap:8px}.hero,.heading{align-items:flex-start;flex-direction:column}.kpis,.grid,.formgrid,.chart-grid,.twocol{grid-template-columns:1fr 1fr}.ratio-grid{grid-template-columns:1fr 1fr}.span2,.span3{grid-column:span 2}.upload{grid-template-columns:1fr}}@media(max-width:430px){.kpis,.grid,.formgrid,.twocol{grid-template-columns:1fr}.span2,.span3{grid-column:span 1}.ratio-grid{grid-template-columns:1fr 1fr}h1{font-size:25px}.wrap{padding:12px 8px 50px}.sec,.hero{padding:16px}}\n'

EMBEDDED_TEMPLATES = {'memo.html': '{% extends "base.html" %}{% block content %}<section class="hero card"><div><span class="eyebrow">CREDIT MEMORANDUM • DRAFT FOR HUMAN REVIEW</span><h1>{{ app.client.name }}</h1><p>{{ result.decision }}</p></div><a class="btn" href="/applications/{{ app.id }}/pdf">تنزيل PDF</a></section><section class="card sec memo"><h2>1. Executive Summary</h2><p>طلب {{ app.facility_type or "تمويل" }} بمبلغ {{ "{:,.0f}".format(app.requested_amount) }} {{ app.currency }} لمدة {{ app.tenor_months|int }} شهر. التقييم {{ result.score|round|int }}/100، ودرجة المخاطر {{ result.grade }} — {{ result.risk }}.</p><h2>2. Borrower Profile</h2><p>{{ app.client.name }} • {{ app.client.sector }} • {{ app.client.legal_form }}.</p><h2>3. Facility Request</h2><p><b>الغرض:</b> {{ app.purpose or "غير مدخل" }}<br><b>مصدر السداد:</b> {{ app.repayment_source or "غير مدخل" }}</p><h2>4. Financial Analysis</h2><p>DSCR = {{ "%.2f"|format(result.dscr) }}x، Current Ratio = {{ "%.2f"|format(result.ratios.current_ratio) }}x، Debt/Equity = {{ "%.2f"|format(result.ratios.debt_equity) }}x، EBITDA Margin = {{ "%.1f"|format(result.ratios.ebitda_margin*100) }}%.</p><h2>5. Key Strengths</h2><ul>{% for x in risk.strengths %}<li>{{ x }}</li>{% endfor %}</ul><h2>6. Key Risks</h2><ul>{% for x in risk.risks %}<li>{{ x }}</li>{% endfor %}</ul><h2>7. Mitigants</h2><ul>{% for x in risk.mitigants %}<li>{{ x }}</li>{% endfor %}</ul><h2>8. Conditions Precedent</h2><ul>{% for x in risk.conditions %}<li>{{ x }}</li>{% endfor %}</ul><h2>9. Credit Committee Questions</h2><ul>{% for x in risk.questions %}<li>{{ x }}</li>{% endfor %}</ul><h2>10. Recommendation</h2><div class="decision {{ color }}">{{ result.decision }}</div><p class="disclaimer">مسودة دعم قرار قابلة للمراجعة البشرية وليست موافقة آلية.</p></section>{% endblock %}\n', 'analysis.html': '{% extends "base.html" %}{% block content %}<section class="hero card"><div><span class="eyebrow">EXPLAINABLE CREDIT ENGINE</span><h1>{{ app.client.name }}</h1><p>تحليل مالي، اتجاهات، Score، مخاطر، ضغط وهيكلة تمويل.</p></div><div class="decision {{ color }}">{{ result.decision }}</div></section><div class="kpis"><div class="kpi"><span>Credit Score</span><b>{{ "%.0f"|format(result.score) }}/100</b></div><div class="kpi"><span>Risk Grade</span><b>{{ result.grade }}</b><small>{{ result.risk }}</small></div><div class="kpi"><span>DSCR</span><b>{{ "%.2f"|format(result.dscr) }}x</b></div><div class="kpi"><span>Indicative Limit</span><b>{{ "{:,.0f}".format(result.indicative_limit) }}</b></div></div><section class="card sec"><h2>20+ Financial Ratios</h2><div class="ratio-grid">{% for name,val in ratio_labels %}<div class="ratio"><span>{{ name }}</span><b>{{ val }}</b></div>{% endfor %}</div></section><section class="card sec"><h2>Trend Analysis</h2><div class="chart-grid">{{ charts|safe }}</div></section><section class="card sec"><div class="twocol"><div><h2>Risk Analysis</h2><h3>نقاط القوة</h3><ul>{% for x in risk.strengths %}<li>{{ x }}</li>{% endfor %}</ul><h3>المخففات</h3><ul>{% for x in risk.mitigants %}<li>{{ x }}</li>{% endfor %}</ul></div><div><h2>المخاطر</h2><ul>{% for x in risk.risks %}<li>{{ x }}</li>{% endfor %}</ul><h3>الشروط</h3><ul>{% for x in risk.conditions %}<li>{{ x }}</li>{% endfor %}</ul></div></div></section><section class="card sec"><h2>Stress Testing</h2><div class="table"><table><tr><th>السيناريو</th><th>DSCR</th><th>الحالة</th></tr>{% for s in result.stresses %}<tr><td>{{ s.name }}</td><td>{{ "%.2f"|format(s.dscr) }}x</td><td>{{ "مقبول" if s.dscr>=1.25 else "حساس" if s.dscr>=1 else "غير مغطى" }}</td></tr>{% endfor %}</table></div></section><section class="card sec"><h2>هيكلة التسهيل</h2><div class="kpis"><div class="kpi"><span>قيد التدفق</span><b>{{ "{:,.0f}".format(result.cash_limit) }}</b></div><div class="kpi"><span>قيد الرافعة</span><b>{{ "{:,.0f}".format(result.leverage_limit) }}</b></div><div class="kpi"><span>قيد الضمان</span><b>{{ "{:,.0f}".format(result.collateral_limit) }}</b></div><div class="kpi"><span>الحد الإرشادي</span><b>{{ "{:,.0f}".format(result.indicative_limit) }}</b></div></div></section>{% endblock %}\n', 'login.html': '{% extends "base.html" %}{% block content %}<div class="auth card"><h1>تسجيل الدخول</h1><form method="post"><label>البريد الإلكتروني</label><input name="email" type="email" required><label>كلمة المرور</label><input name="password" type="password" required><button>دخول</button></form><p>ليس لديك حساب؟ <a href="/register">إنشاء مستخدم</a></p></div>{% endblock %}\n', 'base.html': '<!doctype html><html lang="ar" dir="rtl"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1"><title>{{ title or "AI Credit Analysis" }}</title><link rel="stylesheet" href="/static/style.css"></head><body><header><div class="top"><div><b>AI Credit Analysis</b><small>MVP 1.0 • Corporate Credit Underwriting</small></div>{% if user %}<nav><a href="/dashboard">لوحة التحكم</a><a href="/clients">العملاء</a><a href="/logout">خروج</a></nav>{% endif %}</div></header><main class="wrap">{% if message %}<div class="notice">{{ message }}</div>{% endif %}{% block content %}{% endblock %}</main></body></html>\n', 'register.html': '{% extends "base.html" %}{% block content %}<div class="auth card"><h1>إنشاء مستخدم</h1><form method="post"><label>الاسم</label><input name="full_name" required><label>البريد</label><input name="email" type="email" required><label>كلمة المرور</label><input name="password" type="password" minlength="8" required><button>إنشاء الحساب</button></form></div>{% endblock %}\n', 'client.html': '{% extends "base.html" %}{% block content %}<section class="hero card"><div><span class="eyebrow">BORROWER PROFILE</span><h1>{{ client.name }}</h1><p>{{ client.sector }} • {{ client.legal_form }}</p></div><a class="btn" href="/clients/{{ client.id }}/applications/new">+ طلب تمويل</a></section><section class="card sec"><h2>طلبات التمويل</h2><div class="table"><table><tr><th>الطلب</th><th>النوع</th><th>المبلغ</th><th>الحالة</th><th></th></tr>{% for a in client.applications %}<tr><td>{{ a.title }}</td><td>{{ a.facility_type }}</td><td>{{ "{:,.0f}".format(a.requested_amount) }} {{ a.currency }}</td><td>{{ a.status }}</td><td><a href="/applications/{{ a.id }}">فتح</a></td></tr>{% endfor %}</table></div></section>{% endblock %}\n', 'application.html': '{% extends "base.html" %}{% block content %}<section class="hero card"><div><span class="eyebrow">CREDIT APPLICATION #{{ app.id }}</span><h1>{{ app.client.name }}</h1><p>{{ app.title }} • {{ "{:,.0f}".format(app.requested_amount) }} {{ app.currency }}</p></div><span class="badge">{{ app.status }}</span></section><div class="tabs"><a href="#docs">المستندات</a><a href="/applications/{{ app.id }}/verify">التحقق</a><a href="/applications/{{ app.id }}/analysis">التحليل</a><a href="/applications/{{ app.id }}/memo">Credit Memo</a><a href="/applications/{{ app.id }}/pdf">PDF</a></div><section id="docs" class="card sec"><h2>Document Center</h2><form method="post" action="/applications/{{ app.id }}/upload" enctype="multipart/form-data" class="upload"><input type="file" name="files" multiple required accept=".pdf,.xlsx,.xlsm,.csv,.docx,.txt"><button>رفع واستخراج</button></form><div class="table"><table><tr><th>الملف</th><th>الحجم</th><th>الحالة</th><th>القيم المستخرجة</th></tr>{% for d in app.documents %}<tr><td>{{ d.filename }}</td><td>{{ "%.1f"|format(d.size_bytes/1024) }} KB</td><td>{{ d.extraction_status }}</td><td>{{ d.extracted_count }}</td></tr>{% endfor %}</table></div></section><section class="card sec"><h2>حالة البيانات</h2><div class="kpis"><div class="kpi"><span>القيم</span><b>{{ value_count }}</b></div><div class="kpi"><span>تم التحقق</span><b>{{ verified_count }}</b></div><div class="kpi"><span>غير معتمدة</span><b>{{ value_count-verified_count }}</b></div><div class="kpi"><span>تكلفة API</span><b>${{ "%.4f"|format(api_cost) }}</b></div></div></section>{% endblock %}\n', 'verify.html': '{% extends "base.html" %}{% block content %}<section class="heading"><div><h1>التحقق من البيانات</h1><p>{{ app.client.name }} — لا تعتمد أرقام التحليل قبل مراجعتها.</p></div><a class="btn ghost" href="/applications/{{ app.id }}">رجوع</a></section><section class="card sec"><form method="post"><div class="table"><table><tr><th>اعتماد</th><th>الفترة</th><th>البند</th><th>القيمة</th><th>الثقة</th><th>المصدر</th><th>ملاحظة</th></tr>{% for v in values %}<tr><td><input class="check" type="checkbox" name="verified_{{ v.id }}" {% if v.verified %}checked{% endif %}></td><td>{{ v.period }}</td><td>{{ v.metric }}</td><td><input name="value_{{ v.id }}" value="{{ v.value }}"></td><td>{{ "%.0f"|format(v.confidence*100) }}%</td><td>{{ v.source_location }}</td><td><input name="note_{{ v.id }}" value="{{ v.reviewer_note or \'\' }}"></td></tr>{% endfor %}</table></div><button>حفظ التحقق</button></form></section>{% endblock %}\n', 'dashboard.html': '{% extends "base.html" %}{% block content %}<section class="hero card"><div><span class="eyebrow">MVP 1.0</span><h1>لوحة الائتمان</h1><p>من العميل إلى المذكرة الائتمانية في مسار واحد قابل للتدقيق.</p></div><a class="btn" href="/clients/new">+ عميل جديد</a></section><div class="kpis"><div class="kpi"><span>العملاء</span><b>{{ stats.clients }}</b></div><div class="kpi"><span>طلبات التمويل</span><b>{{ stats.apps }}</b></div><div class="kpi"><span>ملفات مرفوعة</span><b>{{ stats.docs }}</b></div><div class="kpi"><span>تكلفة API</span><b>${{ "%.4f"|format(stats.cost) }}</b></div></div><section class="card sec"><h2>أحدث الطلبات</h2><div class="table"><table><tr><th>العميل</th><th>الطلب</th><th>المبلغ</th><th>الحالة</th><th></th></tr>{% for a in apps %}<tr><td>{{ a.client.name }}</td><td>{{ a.title }}</td><td>{{ "{:,.0f}".format(a.requested_amount) }} {{ a.currency }}</td><td><span class="badge">{{ a.status }}</span></td><td><a href="/applications/{{ a.id }}">فتح</a></td></tr>{% endfor %}</table></div></section>{% endblock %}\n', 'app_new.html': '{% extends "base.html" %}{% block content %}<section class="card sec"><h1>طلب تمويل جديد — {{ client.name }}</h1><form method="post" class="formgrid"><div><label>عنوان الطلب</label><input name="title" value="طلب تمويل"></div><div><label>نوع التسهيل</label><input name="facility_type"></div><div><label>العملة</label><input name="currency" value="YER"></div><div><label>المبلغ المطلوب</label><input name="requested_amount" type="number" step="any"></div><div><label>المدة / شهر</label><input name="tenor_months" type="number" value="12"></div><div><label>العائد السنوي %</label><input name="annual_rate" type="number" step="any"></div><div><label>خدمة الدين السنوية القائمة</label><input name="existing_annual_debt_service" type="number" step="any"></div><div><label>قيمة الضمان</label><input name="collateral_value" type="number" step="any"></div><div><label>التقييم النوعي 0-100</label><input name="qualitative_score" type="number" value="50"></div><div class="span3"><label>الغرض</label><textarea name="purpose"></textarea></div><div class="span3"><label>مصدر السداد</label><textarea name="repayment_source"></textarea></div><button>إنشاء الطلب</button></form></section>{% endblock %}\n', 'client_new.html': '{% extends "base.html" %}{% block content %}<section class="card sec"><h1>إنشاء عميل</h1><form method="post" class="formgrid"><div><label>اسم العميل/الشركة</label><input name="name" required></div><div><label>الشكل القانوني</label><input name="legal_form"></div><div><label>القطاع</label><input name="sector"></div><div><label>رقم السجل</label><input name="registration_no"></div><div><label>المدينة</label><input name="city"></div><div class="span2"><label>ملاحظات</label><textarea name="notes"></textarea></div><button>حفظ العميل</button></form></section>{% endblock %}\n', 'clients.html': '{% extends "base.html" %}{% block content %}<section class="heading"><div><h1>العملاء</h1><p>ملفات المقترضين وطلبات التمويل.</p></div><a class="btn" href="/clients/new">+ عميل جديد</a></section><div class="grid">{% for c in clients %}<a class="client card" href="/clients/{{ c.id }}"><b>{{ c.name }}</b><span>{{ c.sector or "قطاع غير محدد" }}</span><small>{{ c.legal_form }}</small></a>{% endfor %}</div>{% endblock %}\n'}


def ensure_embedded_assets():
    base=Path('app')
    (base/'templates').mkdir(parents=True,exist_ok=True)
    (base/'static').mkdir(parents=True,exist_ok=True)
    (base/'static'/'style.css').write_text(EMBEDDED_CSS,encoding='utf-8')
    for name,content in EMBEDDED_TEMPLATES.items():
        (base/'templates'/name).write_text(content,encoding='utf-8')
ensure_embedded_assets()

from fastapi import FastAPI,Request,Form,UploadFile,File,Depends,HTTPException
from fastapi.responses import HTMLResponse,RedirectResponse,Response
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from starlette.middleware.sessions import SessionMiddleware
from sqlalchemy.orm import Session
from sqlalchemy import func
from typing import List
from pathlib import Path
import os,uuid,html
Base.metadata.create_all(bind=engine)
app=FastAPI(title='AI Credit Analysis MVP 1.0')
app.add_middleware(SessionMiddleware,secret_key=os.getenv('APP_SECRET','change-me-before-production'),same_site='lax',https_only=False)
app.mount('/static',StaticFiles(directory='app/static'),name='static');templates=Jinja2Templates(directory='app/templates')
UPLOAD_DIR=Path(os.getenv('UPLOAD_DIR','./uploads'));UPLOAD_DIR.mkdir(parents=True,exist_ok=True);MAX_FILE_MB=int(os.getenv('MAX_FILE_MB','20'));MAX_TOTAL_MB=int(os.getenv('MAX_TOTAL_MB','60'))
def current_user(request,db):
 uid=request.session.get('user_id');return db.get(User,uid) if uid else None
def require_user(request,db):
 u=current_user(request,db)
 if not u:raise HTTPException(status_code=401)
 return u
def own_client(db,user,cid):
 c=db.query(Client).filter(Client.id==cid,Client.user_id==user.id).first()
 if not c:raise HTTPException(404)
 return c
def own_app(db,user,aid):
 a=db.query(CreditApplication).join(Client).filter(CreditApplication.id==aid,Client.user_id==user.id).first()
 if not a:raise HTTPException(404)
 return a
def render(name,request,**ctx):return templates.TemplateResponse(name,{'request':request,**ctx})
@app.get('/')
def index(request:Request,db:Session=Depends(get_db)):return RedirectResponse('/dashboard' if current_user(request,db) else '/login',303)
@app.get('/register',response_class=HTMLResponse)
def register_get(request:Request):return render('register.html',request,user=None)
@app.post('/register')
def register_post(request:Request,full_name:str=Form(...),email:str=Form(...),password:str=Form(...),db:Session=Depends(get_db)):
 if db.query(User).filter(func.lower(User.email)==email.lower()).first():return render('register.html',request,user=None,message='البريد مستخدم بالفعل.')
 u=User(full_name=full_name,email=email.lower(),password_hash=hash_password(password));db.add(u);db.commit();db.refresh(u);request.session['user_id']=u.id;return RedirectResponse('/dashboard',303)
@app.get('/login',response_class=HTMLResponse)
def login_get(request:Request):return render('login.html',request,user=None)
@app.post('/login')
def login_post(request:Request,email:str=Form(...),password:str=Form(...),db:Session=Depends(get_db)):
 u=db.query(User).filter(func.lower(User.email)==email.lower()).first()
 if not u or not verify_password(password,u.password_hash):return render('login.html',request,user=None,message='بيانات الدخول غير صحيحة.')
 request.session['user_id']=u.id;return RedirectResponse('/dashboard',303)
@app.get('/logout')
def logout(request:Request):request.session.clear();return RedirectResponse('/login',303)
@app.get('/dashboard',response_class=HTMLResponse)
def dashboard(request:Request,db:Session=Depends(get_db)):
 u=require_user(request,db);apps=db.query(CreditApplication).join(Client).filter(Client.user_id==u.id).order_by(CreditApplication.created_at.desc()).limit(10).all();stats={'clients':db.query(Client).filter(Client.user_id==u.id).count(),'apps':db.query(CreditApplication).join(Client).filter(Client.user_id==u.id).count(),'docs':db.query(Document).join(CreditApplication).join(Client).filter(Client.user_id==u.id).count(),'cost':db.query(func.coalesce(func.sum(ApiCost.cost_usd),0)).join(CreditApplication,ApiCost.application_id==CreditApplication.id).join(Client).filter(Client.user_id==u.id).scalar() or 0};return render('dashboard.html',request,user=u,stats=stats,apps=apps)
@app.get('/clients',response_class=HTMLResponse)
def clients(request:Request,db:Session=Depends(get_db)):
 u=require_user(request,db);return render('clients.html',request,user=u,clients=db.query(Client).filter(Client.user_id==u.id).order_by(Client.created_at.desc()).all())
@app.get('/clients/new',response_class=HTMLResponse)
def client_new_get(request:Request,db:Session=Depends(get_db)):return render('client_new.html',request,user=require_user(request,db))
@app.post('/clients/new')
def client_new_post(request:Request,name:str=Form(...),legal_form:str=Form(''),sector:str=Form(''),registration_no:str=Form(''),city:str=Form(''),notes:str=Form(''),db:Session=Depends(get_db)):
 u=require_user(request,db);c=Client(user_id=u.id,name=name,legal_form=legal_form,sector=sector,registration_no=registration_no,city=city,notes=notes);db.add(c);db.commit();db.refresh(c);return RedirectResponse(f'/clients/{c.id}',303)
@app.get('/clients/{cid}',response_class=HTMLResponse)
def client_view(cid:int,request:Request,db:Session=Depends(get_db)):
 u=require_user(request,db);return render('client.html',request,user=u,client=own_client(db,u,cid))
@app.get('/clients/{cid}/applications/new',response_class=HTMLResponse)
def app_new_get(cid:int,request:Request,db:Session=Depends(get_db)):
 u=require_user(request,db);return render('app_new.html',request,user=u,client=own_client(db,u,cid))
@app.post('/clients/{cid}/applications/new')
def app_new_post(cid:int,request:Request,title:str=Form('طلب تمويل'),facility_type:str=Form(''),currency:str=Form('YER'),requested_amount:float=Form(0),tenor_months:float=Form(12),annual_rate:float=Form(0),existing_annual_debt_service:float=Form(0),collateral_value:float=Form(0),qualitative_score:float=Form(50),purpose:str=Form(''),repayment_source:str=Form(''),db:Session=Depends(get_db)):
 u=require_user(request,db);c=own_client(db,u,cid);a=CreditApplication(client_id=c.id,title=title,facility_type=facility_type,currency=currency,requested_amount=requested_amount,tenor_months=tenor_months,annual_rate=annual_rate,existing_annual_debt_service=existing_annual_debt_service,collateral_value=collateral_value,qualitative_score=qualitative_score,purpose=purpose,repayment_source=repayment_source);db.add(a);db.commit();db.refresh(a);return RedirectResponse(f'/applications/{a.id}',303)
@app.get('/applications/{aid}',response_class=HTMLResponse)
def app_view(aid:int,request:Request,db:Session=Depends(get_db)):
 u=require_user(request,db);a=own_app(db,u,aid);vc=db.query(FinancialValue).filter(FinancialValue.application_id==aid).count();ver=db.query(FinancialValue).filter(FinancialValue.application_id==aid,FinancialValue.verified==True).count();cost=db.query(func.coalesce(func.sum(ApiCost.cost_usd),0)).filter(ApiCost.application_id==aid).scalar() or 0;return render('application.html',request,user=u,app=a,value_count=vc,verified_count=ver,api_cost=cost)
@app.post('/applications/{aid}/upload')
async def upload(aid:int,request:Request,files:List[UploadFile]=File(...),db:Session=Depends(get_db)):
 u=require_user(request,db);a=own_app(db,u,aid);total=0
 for f in files[:30]:
  data=await f.read();total+=len(data)
  if len(data)>MAX_FILE_MB*1024*1024 or total>MAX_TOTAL_MB*1024*1024:raise HTTPException(413,'File limit exceeded')
  safe=f"{uuid.uuid4().hex}_{Path(f.filename or 'document').name}";path=UPLOAD_DIR/safe;path.write_bytes(data);d=Document(application_id=aid,filename=f.filename or safe,stored_path=str(path),file_type=Path(safe).suffix.lower(),size_bytes=len(data),extraction_status='processing');db.add(d);db.commit();db.refresh(d)
  try:
   vals=extract_file(d.filename,data)
   for x in vals:db.add(FinancialValue(application_id=aid,document_id=d.id,period=x['period'],metric=x['metric'],value=x['value'],source_location=x['location'],confidence=x['confidence'],verified=False))
   d.extracted_count=len(vals);d.extraction_status='extracted' if vals else 'review_required'
  except Exception:d.extraction_status='failed'
  db.commit()
 a.status='data_extracted';db.commit();return RedirectResponse(f'/applications/{aid}/verify',303)
@app.get('/applications/{aid}/verify',response_class=HTMLResponse)
def verify_get(aid:int,request:Request,db:Session=Depends(get_db)):
 u=require_user(request,db);a=own_app(db,u,aid);vals=db.query(FinancialValue).filter(FinancialValue.application_id==aid).order_by(FinancialValue.period,FinancialValue.metric).all();return render('verify.html',request,user=u,app=a,values=vals)
@app.post('/applications/{aid}/verify')
async def verify_post(aid:int,request:Request,db:Session=Depends(get_db)):
 u=require_user(request,db);a=own_app(db,u,aid);form=await request.form();vals=db.query(FinancialValue).filter(FinancialValue.application_id==aid).all()
 for v in vals:
  v.verified=f'verified_{v.id}' in form
  try:v.value=float(str(form.get(f'value_{v.id}',v.value)).replace(',',''))
  except:pass
  v.reviewer_note=str(form.get(f'note_{v.id}',''))
 a.status='verified';db.commit();return RedirectResponse(f'/applications/{aid}/analysis',303)
def make_chart(periods,key,title):
 vals=[float(x.get(key,0) or 0) for x in periods];labs=[x['period'] for x in periods]
 if not vals:return f'<div class="chart"><b>{html.escape(title)}</b><p>لا بيانات</p></div>'
 W,H,P=620,220,42;mn=min(vals);mx=max(vals);mx=mn+1 if mx==mn else mx;pts=[]
 for i,v in enumerate(vals):
  x=P+(W-2*P)*i/max(len(vals)-1,1);y=H-P-(H-2*P)*(v-mn)/(mx-mn);pts.append((x,y,v))
 path=' '.join(('M' if i==0 else 'L')+f' {x:.1f} {y:.1f}' for i,(x,y,v) in enumerate(pts));dots=''.join(f'<circle cx="{x}" cy="{y}" r="5"/><text x="{x}" y="{max(16,y-9)}" text-anchor="middle">{v:,.0f}</text><text x="{x}" y="{H-10}" text-anchor="middle">{html.escape(labs[i])}</text>' for i,(x,y,v) in enumerate(pts));return f'<div class="chart"><b>{html.escape(title)}</b><svg viewBox="0 0 {W} {H}"><path d="{path}"/>{dots}</svg></div>'
def analyze(db,a):
 vals=db.query(FinancialValue).filter(FinancialValue.application_id==a.id).all();periods=periods_from_values(vals,verified_only=False);result=score_application(periods,a);risk=local_risk_analysis(periods,a,result);return periods,result,risk
@app.get('/applications/{aid}/analysis',response_class=HTMLResponse)
def analysis(aid:int,request:Request,db:Session=Depends(get_db)):
 u=require_user(request,db);a=own_app(db,u,aid);periods,result,risk=analyze(db,a);r=result['ratios'];labels=[('Gross Margin',f"{r['gross_margin']*100:.1f}%"),('EBITDA Margin',f"{r['ebitda_margin']*100:.1f}%"),('Operating Margin',f"{r['operating_margin']*100:.1f}%"),('Net Margin',f"{r['net_margin']*100:.1f}%"),('ROA',f"{r['roa']*100:.1f}%"),('ROE',f"{r['roe']*100:.1f}%"),('Asset Turnover',f"{r['asset_turnover']:.2f}x"),('Current Ratio',f"{r['current_ratio']:.2f}x"),('Quick Ratio',f"{r['quick_ratio']:.2f}x"),('Cash Ratio',f"{r['cash_ratio']:.2f}x"),('Working Capital',f"{r['working_capital']:,.0f}"),('Debt/Equity',f"{r['debt_equity']:.2f}x"),('Debt/Assets',f"{r['debt_assets']*100:.1f}%"),('Liabilities/Assets',f"{r['liabilities_assets']*100:.1f}%"),('Equity Ratio',f"{r['equity_ratio']*100:.1f}%"),('Interest Coverage',f"{r['interest_coverage']:.2f}x"),('CFO Margin',f"{r['cfo_margin']*100:.1f}%"),('CFO/Debt',f"{r['cfo_debt']*100:.1f}%"),('Free Cash Flow',f"{r['free_cash_flow']:,.0f}"),('DSO',f"{r['dso']:.0f} d"),('DIO',f"{r['dio']:.0f} d"),('DPO',f"{r['dpo']:.0f} d"),('Cash Conversion Cycle',f"{r['cash_conversion_cycle']:.0f} d"),('Inventory Turnover',f"{r['inventory_turnover']:.2f}x"),('Receivables Turnover',f"{r['receivables_turnover']:.2f}x"),('Payables Turnover',f"{r['payables_turnover']:.2f}x")];charts=''.join([make_chart(periods,'revenue','الإيرادات'),make_chart(periods,'ebitda','EBITDA'),make_chart(periods,'cfo','التدفق التشغيلي'),make_chart(periods,'total_debt','إجمالي الدين')]);color='green' if 'مؤهل' in result['decision'] else 'red' if 'لا يوصى' in result['decision'] else 'amber';return render('analysis.html',request,user=u,app=a,result=result,risk=risk,ratio_labels=labels,charts=charts,color=color)
@app.get('/applications/{aid}/memo',response_class=HTMLResponse)
def memo(aid:int,request:Request,db:Session=Depends(get_db)):
 u=require_user(request,db);a=own_app(db,u,aid);periods,result,risk=analyze(db,a);color='green' if 'مؤهل' in result['decision'] else 'red' if 'لا يوصى' in result['decision'] else 'amber';return render('memo.html',request,user=u,app=a,result=result,risk=risk,color=color)
@app.get('/applications/{aid}/pdf')
def pdf(aid:int,request:Request,db:Session=Depends(get_db)):
 u=require_user(request,db);a=own_app(db,u,aid);periods,result,risk=analyze(db,a);data=build_pdf(a.client,a,periods,result,risk);return Response(data,media_type='application/pdf',headers={'Content-Disposition':f'attachment; filename="credit_memo_{aid}.pdf"'})
@app.get('/health')
def health():return {'status':'ok','version':'1.0.0-mvp','features':['auth','clients','applications','documents','extraction','verification','ratios','trends','score','risk_analysis','credit_memo','pdf','api_cost','dashboard']}
